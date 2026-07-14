"""
Convert Chuong4.md to Chuong4.docx with proper formatting.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def add_run(paragraph, text, bold=False, italic=False, font_name='Times New Roman', font_size=13):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def set_cell_text(cell, text, bold=False, font_size=11):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)

def convert_simple_table(md_table, doc):
    """Convert a markdown table to a docx table."""
    lines = [l.strip() for l in md_table.strip().split('\n') if l.strip()]
    if not lines:
        return
    # Skip separator line (|---|---|)
    data_lines = [l for l in lines if not re.match(r'^[\s\|:\-]+$', l)]
    if not data_lines:
        return
    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split('|')]
        # Remove empty first/last from leading/trailing |
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                is_header = (i == 0)
                set_cell_text(cell, cell_text, bold=is_header,
                              font_size=11 if is_header else 10)

def md_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Process line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith('```'):
            if in_code_block:
                # Write code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                code_text = '\n'.join(code_buffer)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(50, 50, 50)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            # Flush table if any
            if table_buffer:
                convert_simple_table('\n'.join(table_buffer), doc)
                table_buffer = []
                in_table = False
                doc.add_paragraph()  # spacing
            i += 1
            continue

        # Table detection
        if line.strip().startswith('|') and '|' in line:
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        else:
            if table_buffer:
                convert_simple_table('\n'.join(table_buffer), doc)
                table_buffer = []
                in_table = False
                doc.add_paragraph()

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(line[5:])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        # Bullet list
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            # Handle bold within list items
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    add_run(p, part[2:-2], bold=True)
                else:
                    add_run(p, part)
            p.paragraph_format.space_after = Pt(2)
        # Numbered list
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(2)
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    add_run(p, part[2:-2], bold=True)
                else:
                    add_run(p, part)
        # Regular paragraph
        else:
            text = line.strip()
            # Skip standalone formatting like ---
            if re.match(r'^[-=]{3,}$', text):
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.3
            # Handle bold and code
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    add_run(p, part[2:-2], bold=True)
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    add_run(p, part)

        i += 1

    # Flush remaining table
    if table_buffer:
        convert_simple_table('\n'.join(table_buffer), doc)

    doc.save(docx_file)
    print(f"Da xuat file: {docx_file}")

if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    md_path = os.path.join(os.path.dirname(__file__), 'Chuong4.md')
    docx_path = os.path.join('C:\\Users\\atheng\\AppData\\Local\\Temp', 'Chuong4.docx')
    md_to_docx(md_path, docx_path)
